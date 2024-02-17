import tkinter as tk
from tkinter import ttk, messagebox
import tkinter.filedialog as filedialog
import sqlite3
from docx import Document


class CGPACalculatorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("GradePoint Pro")
        self.root.geometry("1400x900")
        self.create_widgets()
        self.style = ttk.Style()
        self.style.theme_use("clam")
        self.style.configure(
            "TButton",
            foreground="white",
            background="#3498db",
            font=("Helvetica", 10, "bold"),
            padding=10,
        )
        self.style.map("TButton", background=[("active", "#2980b9")])
        self.style.configure(
            "Title.TLabel",
            foreground="#3498db",
            font=("Helvetica", 20, "bold"),
        )

        self.style.configure(
            "Instructions.TFrame",
            background="#f0f0f0",
            padding=10,
            borderwidth=2,
            relief="groove",
            font=("Helvetica", 14),
        )

        # Connect to SQLite3 database
        self.conn = sqlite3.connect("gradepointpro.db")
        self.cursor = self.conn.cursor()
        self.create_tables()

    def create_tables(self):
        """
        Create the 'unweighted' and 'weighted' tables in the database if they don't exist.
        """
        self.cursor.execute(
            """CREATE TABLE IF NOT EXISTS unweighted (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            course_name TEXT,
                            credits REAL,
                            grade TEXT)"""
        )
        self.cursor.execute(
            """CREATE TABLE IF NOT EXISTS weighted (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            course_type TEXT,
                            course_name TEXT,
                            credits REAL,
                            grade TEXT)"""
        )
        self.conn.commit()

    def insert_unweighted_course(self, course_name, credits, grade):
        """
        Insert a new unweighted course entry into the 'unweighted' table.
        """
        self.cursor.execute(
            "INSERT INTO unweighted (course_name, credits, grade) VALUES (?, ?, ?)",
            (course_name, credits, grade),
        )
        self.conn.commit()

    def insert_weighted_course(self, course_type, course_name, credits, grade):
        """
        Insert a new weighted course entry into the 'weighted' table.
        """
        self.cursor.execute(
            "INSERT INTO weighted (course_type, course_name, credits, grade) VALUES (?, ?, ?, ?)",
            (course_type, course_name, credits, grade),
        )
        self.conn.commit()

    def create_widgets(self):
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=100)

        self.instructions_frame = ttk.Frame(self.notebook, style="Instructions.TFrame")
        self.notebook.add(self.instructions_frame, text="Instructions")
        self.create_instructions_widgets()

        self.unweighted_frame = ttk.Frame(self.notebook, style="Instructions.TFrame")
        self.notebook.add(self.unweighted_frame, text="Unweighted GPA Calculator")
        self.create_unweighted_widgets()

        self.weighted_frame = ttk.Frame(self.notebook, style="Instructions.TFrame")
        self.notebook.add(self.weighted_frame, text="Weighted GPA Calculator")
        self.create_weighted_widgets()

        self.faq_frame = ttk.Frame(self.notebook, style="Instructions.TFrame")
        self.notebook.add(self.faq_frame, text="FAQ")
        self.create_faq_widgets()

        self.notebook.enable_traversal()

    def create_instructions_widgets(self):
        instructions_text = """
        Welcome to GradePointPro's GPA Calculator

        To run the application, open the file in your preferred IDE (recommended through vscode) and run the program.

        To receive your exact GPA, follow these steps:

        1. Choose whether you want to calculate your Unweighted or Weighted GPA.
        2. After selecting the correct calculator:
           - For your unweighted GPA, input the course name, number of credits, and your current grade in the class.
           - For the weighted calculator, select the course type (Regular, Honors, AP, AICE, or Dual Enrollment), add the course name, include the number of credits, and your current grade.
        3. Click on the '+' button to add the course. You can remove a course by clicking the '-' button.
        4. Repeat step 2 until you have all your desired courses added.
        5. Click on 'Calculate GPA' to view your GPA based on the added courses. You can also add any class that you plan to take in the future.
        """

        instructions_label = tk.Label(
            self.instructions_frame,
            text=instructions_text,
            font=("Calibri", 14, "bold"),
            justify="center",
            padx=20,
            pady=20,
            wraplength=800,
        )
        instructions_label.pack(fill="both", expand=True)

    def create_faq_widgets(self):
        faq_data = [
            (
                "Can people from other schools use this?",
                "The gpa calculator is free to use for anyone from any school. Our calculators use the Hillsborough County Public Schools GPA scale to calculate. Before using the GPA calculator we reccommend that users review their school’s or district’s GPA scale to make sure that the calculations match.",
            ),
            (
                "What Grading Scale does the GPA Calculator use?",
                "The grading scale the the calculators use is from the Hillsborough County Public School’s given grading/GPAchart.",
            ),
            (
                "What is the difference between the Unweighted GPA and Weighted GPA?",
                "When calculating your weighted GPA the calculator takes the credits you get for a class into you GPA, generally making it a higher number than the unweighted gpa. The unweighted GPA calculator only takes your class grade into account, in which you can get a maximum of 4.0 .",
            ),
            (
                "What is the purpose of the Course Type Menu",
                "The Course Type Menu is used to determine which level of class that you’re taking. Based on the level of the class, a certain number of credits is added onto your  weighted GPA. For an regular class there is no extra credit, for an honors class an extra 0.04 of credits is added on to you unweighted GPA, while a AP, AICE, or Dual Enrollment class adds 0.08 credits.",
            ),
            (
                "Can you use this calculator to predict future semesters?",
                "You are able to predict or get an estimate of your GPA for the future. To do this use the calculator normally and input what you think  you will get in a class to find out you GPA. You can also use this method to determine whether or not to take a specific class.",
            ),
        ]

        for idx, (question, answer) in enumerate(faq_data):
            question_frame = ttk.Frame(self.faq_frame)
            question_frame.pack(padx=20, pady=10, fill="both", expand=True)

            question_label = tk.Label(
                question_frame,
                text=f"{idx + 1}. {question}",
                font=("Helvetica", 14, "bold"),
                padx=10,
                pady=10,
                wraplength=800,
                cursor="hand2",
            )
            question_label.pack(anchor="w", fill="x")
            question_label.bind(
                "<Button-1>",
                lambda event, frame=question_frame, answer=answer: self.toggle_answer_visibility(
                    frame, answer
                ),
            )

    def toggle_answer_visibility(self, frame, answer):
        children = frame.winfo_children()
        if len(children) == 1:
            answer_label = tk.Label(
                frame,
                text=answer,
                font=("Helvetica", 12),
                padx=10,
                pady=10,
                wraplength=800,
            )
            answer_label.pack(anchor="w", fill="x")
            frame.update_idletasks()
        else:
            answer_label = children[1]
            if answer_label.winfo_ismapped():
                answer_label.pack_forget()
            else:
                answer_label.pack(anchor="w", fill="x")

    def create_unweighted_widgets(self):
        unweighted_box = ttk.Frame(self.unweighted_frame, style="Instructions.TFrame")
        unweighted_box.pack(padx=20, pady=20, fill="both", expand=True)

        ttk.Label(
            unweighted_box,
            text="Unweighted GPA Calculator",
            style="Title.TLabel",
        ).pack(pady=10)

        unweighted_calculator_frame = ttk.Frame(unweighted_box)
        unweighted_calculator_frame.pack(padx=20, pady=20)

        # Labels for course name, credits, and grade
        ttk.Label(unweighted_calculator_frame, text="Course Name").grid(
            row=0, column=0, padx=5, pady=5, sticky="w"
        )
        ttk.Label(unweighted_calculator_frame, text="Credits").grid(
            row=0, column=1, padx=5, pady=5, sticky="w"
        )
        ttk.Label(unweighted_calculator_frame, text="Grade").grid(
            row=0, column=2, padx=5, pady=5, sticky="w"
        )

        self.unweighted_course_wrapper = ttk.Frame(unweighted_calculator_frame)
        self.unweighted_course_wrapper.grid(row=1, column=0, columnspan=3)

        self.add_unweighted_course()

        self.unweighted_buttons_frame = ttk.Frame(unweighted_calculator_frame)
        self.unweighted_buttons_frame.grid(row=2, column=0, columnspan=3, pady=10)

        ttk.Button(
            self.unweighted_buttons_frame,
            text="+ Add Course",
            command=self.add_unweighted_course,
            style="TButton",
        ).grid(row=0, column=0, padx=5)
        ttk.Button(
            self.unweighted_buttons_frame,
            text="- Remove Course",
            command=self.remove_unweighted_course,
            style="TButton",
        ).grid(row=0, column=1, padx=5)
        ttk.Button(
            self.unweighted_buttons_frame,
            text="Calculate GPA",
            command=self.calculate_cgpa,
            style="TButton",
        ).grid(row=0, column=2, padx=5)
        ttk.Button(
            self.unweighted_buttons_frame,
            text="Print",
            command=self.print_unweighted,
            style="TButton",
        ).grid(row=0, column=3, padx=5)

        self.cgpa_label = ttk.Label(
            unweighted_calculator_frame,
            text="Your Unweighted GPA is: ",
            font=("Helvetica", 14),
            foreground="#3498db",
        )
        self.cgpa_label.grid(row=3, column=0, columnspan=3, pady=(10, 0))

        # Table to display course details
        self.unweighted_course_details_table = ttk.Treeview(unweighted_calculator_frame)
        self.unweighted_course_details_table["columns"] = ("Credits", "Grade")
        self.unweighted_course_details_table.heading("#0", text="Course Name")
        self.unweighted_course_details_table.heading("Credits", text="Credits")
        self.unweighted_course_details_table.heading("Grade", text="Grade")
        self.unweighted_course_details_table.grid(
            row=4, column=0, columnspan=3, pady=(10, 0)
        )
        self.unweighted_course_details_table.grid_remove()

    def create_weighted_widgets(self):

        weighted_box = ttk.Frame(self.weighted_frame, style="Instructions.TFrame")
        weighted_box.pack(padx=20, pady=20, fill="both", expand=True)

        ttk.Label(
            weighted_box,
            text="Weighted GPA Calculator",
            style="Title.TLabel",
        ).pack(pady=10)

        weighted_calculator_frame = ttk.Frame(weighted_box)
        weighted_calculator_frame.pack(padx=20, pady=20)

        # Labels for course type, course name, credits, and grade
        ttk.Label(weighted_calculator_frame, text="Course Type").grid(
            row=0, column=0, padx=5, pady=5, sticky="w"
        )
        ttk.Label(weighted_calculator_frame, text="Course Name").grid(
            row=0, column=1, padx=5, pady=5, sticky="w"
        )
        ttk.Label(weighted_calculator_frame, text="Credits").grid(
            row=0, column=2, padx=5, pady=5, sticky="w"
        )
        ttk.Label(weighted_calculator_frame, text="Grade").grid(
            row=0, column=3, padx=5, pady=5, sticky="w"
        )

        self.weighted_course_wrapper = ttk.Frame(weighted_calculator_frame)
        self.weighted_course_wrapper.grid(row=1, column=0, columnspan=4)

        self.add_weighted_course()

        self.weighted_buttons_frame = ttk.Frame(weighted_calculator_frame)
        self.weighted_buttons_frame.grid(row=2, column=0, columnspan=4, pady=10)

        ttk.Button(
            self.weighted_buttons_frame,
            text="+ Add Course",
            command=self.add_weighted_course,
            style="TButton",
        ).grid(row=0, column=0, padx=5)
        ttk.Button(
            self.weighted_buttons_frame,
            text="- Remove Course",
            command=self.remove_weighted_course,
            style="TButton",
        ).grid(row=0, column=1, padx=5)
        ttk.Button(
            self.weighted_buttons_frame,
            text="Calculate GPA",
            command=self.calculate_weighted_cgpa,
            style="TButton",
        ).grid(row=0, column=2, padx=5)
        ttk.Button(
            self.weighted_buttons_frame,
            text="Print",
            command=self.print_weighted,
            style="TButton",
        ).grid(row=0, column=3, padx=5)

        self.weighted_cgpa_label = ttk.Label(
            weighted_calculator_frame,
            text="Your Weighted GPA is:",
            font=("Helvetica", 14),
            foreground="#3498db",
        )
        self.weighted_cgpa_label.grid(row=3, column=0, columnspan=4, pady=(10, 0))

        # Table to display course details
        self.weighted_course_details_table = ttk.Treeview(weighted_calculator_frame)
        self.weighted_course_details_table["columns"] = (
            "Course Name",
            "Credits",
            "Grade",
        )
        self.weighted_course_details_table.heading("#0", text="Course Type")
        self.weighted_course_details_table.heading("Course Name", text="Course Name")
        self.weighted_course_details_table.heading("Credits", text="Credits")
        self.weighted_course_details_table.heading("Grade", text="Grade")
        self.weighted_course_details_table.grid(
            row=4, column=0, columnspan=4, pady=(10, 0)
        )
        self.weighted_course_details_table.grid_remove()

    def add_unweighted_course(self):
        course_frame = ttk.Frame(self.unweighted_course_wrapper)
        course_frame.pack(pady=5)

        ttk.Entry(course_frame).grid(row=0, column=0, padx=5, sticky="w")
        ttk.Entry(course_frame).grid(row=0, column=1, padx=5, sticky="w")
        ttk.Combobox(course_frame, values=["A", "B", "C", "D", "F"]).grid(
            row=0, column=2, padx=5, sticky="w"
        )

    def remove_unweighted_course(self):
        children = self.unweighted_course_wrapper.winfo_children()
        if len(children) > 1:
            children[-1].destroy()
        else:
            messagebox.showwarning("Warning", "You must have at least one course.")

    def add_weighted_course(self):
        course_frame = ttk.Frame(self.weighted_course_wrapper)
        course_frame.pack(pady=5)

        ttk.Combobox(
            course_frame, values=["AP", "AICE", "DE", "Honors", "Regular"]
        ).grid(row=0, column=0, padx=5, sticky="w")
        ttk.Entry(course_frame).grid(row=0, column=1, padx=5, sticky="w")
        ttk.Entry(course_frame).grid(row=0, column=2, padx=5, sticky="w")
        ttk.Combobox(course_frame, values=["A", "B", "C", "D", "F"]).grid(
            row=0, column=3, padx=5, sticky="w"
        )

    def remove_weighted_course(self):
        children = self.weighted_course_wrapper.winfo_children()
        if len(children) > 1:
            children[-1].destroy()
        else:
            messagebox.showwarning("Warning", "You must have at least one course.")

    def calculate_cgpa(self):
        children = self.unweighted_course_wrapper.winfo_children()
        total_grade_points = 0
        total_credit_units = 0

        for child in children:
            entries = child.winfo_children()
            course_code = entries[0].get()
            credit_units = entries[1].get()
            grade = entries[2].get()

            if course_code and credit_units and grade:
                grade_points = {"A": 4, "B": 3, "C": 2, "D": 1, "F": 0}
                total_grade_points += grade_points[grade] * float(credit_units)
                total_credit_units += float(credit_units)

                # Insert course into database
                self.insert_unweighted_course(course_code, credit_units, grade)

        if total_credit_units == 0:
            messagebox.showwarning(
                "Warning", "Please add at least one course with credit units."
            )
        else:
            cgpa = total_grade_points / total_credit_units
            self.cgpa_label.config(text=f"Your Unweighted GPA is: {cgpa:.2f}")

            # Populate the table with course details
            self.unweighted_course_details_table.delete(
                *self.unweighted_course_details_table.get_children()
            )
            for child in children:
                entries = child.winfo_children()
                course_name = entries[0].get()
                credit_units = entries[1].get()
                grade = entries[2].get()
                self.unweighted_course_details_table.insert(
                    "", tk.END, text=course_name, values=(credit_units, grade)
                )
            self.unweighted_course_details_table.grid()

    def calculate_weighted_cgpa(self):
        children = self.weighted_course_wrapper.winfo_children()
        total_unweighted_grade_points = 0
        total_credit_units = 0
        ap_count = 0
        honors_count = 0

        for child in children:
            entries = child.winfo_children()
            course_type = entries[0].get()
            course_name = entries[1].get()
            credit_units = float(entries[2].get())
            grade = entries[3].get()

            if course_name and credit_units and grade:
                self.insert_weighted_course(
                    course_type, course_name, credit_units, grade
                )

                grade_points = {"A": 4, "B": 3, "C": 2, "D": 1, "F": 0}
                unweighted_grade_points = grade_points[grade]

                # Calculate the unweighted grade points
                total_unweighted_grade_points += unweighted_grade_points * credit_units
                total_credit_units += credit_units

                # Count the number of AP, AICE, DE, and Honors classes
                if course_type in ["AP", "AICE", "DE"] and grade in ["A", "B", "C"]:
                    ap_count += 1
                elif course_type == "Honors" and grade in ["A", "B", "C"]:
                    honors_count += 1

        if total_credit_units == 0:
            messagebox.showwarning(
                "Warning", "Please add at least one course with credit units."
            )
        else:
            # Calculate the unweighted GPA
            unweighted_gpa = total_unweighted_grade_points / total_credit_units

            # Add additional points for AP, AICE, DE, and Honors classes
            weighted_gpa = unweighted_gpa + 0.08 * ap_count + 0.04 * honors_count

            self.weighted_cgpa_label.config(
                text=f"Your Weighted GPA is: {weighted_gpa:.2f}"
            )

            # Populate the table with course details
            self.weighted_course_details_table.delete(
                *self.weighted_course_details_table.get_children()
            )
            for child in children:
                entries = child.winfo_children()
                course_type = entries[0].get()
                course_name = entries[1].get()
                credit_units = entries[2].get()
                grade = entries[3].get()
                self.weighted_course_details_table.insert(
                    "",
                    tk.END,
                    text=course_type,
                    values=(course_name, credit_units, grade),
                )
            self.weighted_course_details_table.grid()

    def print_unweighted(self):
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
            title="Save Unweighted GPA Details",
        )
        if filename:
            document = Document()
            document.add_heading("Unweighted GPA Details", level=1)

            for child in self.unweighted_course_wrapper.winfo_children():
                course_name, credits, grade = [
                    widget.get() for widget in child.winfo_children()
                ]
                document.add_paragraph(f"Course Name: {course_name}")
                document.add_paragraph(f"Credits: {credits}")
                document.add_paragraph(f"Grade: {grade}")
                document.add_paragraph("----------------------------------")

            document.add_paragraph(self.cgpa_label["text"])

            document.save(filename)
            messagebox.showinfo(
                "File Saved",
                f"Unweighted GPA details saved successfully at {filename}.",
            )

    def print_weighted(self):
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
            title="Save Weighted GPA Details",
        )
        if filename:
            document = Document()
            document.add_heading("Weighted GPA Details", level=1)

            for child in self.weighted_course_wrapper.winfo_children():
                course_type, course_name, credits, grade = [
                    widget.get() for widget in child.winfo_children()
                ]
                document.add_paragraph(f"Course Type: {course_type}")
                document.add_paragraph(f"Course Name: {course_name}")
                document.add_paragraph(f"Credits: {credits}")
                document.add_paragraph(f"Grade: {grade}")
                document.add_paragraph("----------------------------------")

            document.add_paragraph(self.weighted_cgpa_label["text"])

            document.save(filename)
            messagebox.showinfo(
                "File Saved",
                f"Weighted GPA details saved successfully at {filename}.",
            )


root = tk.Tk()
app = CGPACalculatorApp(root)
root.mainloop()
